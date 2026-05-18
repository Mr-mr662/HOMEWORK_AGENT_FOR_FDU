"""
文档生成器 - 生成Word、PDF等格式的作业文件
"""
from typing import Optional, Dict, Any
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..platforms.base import Assignment


class DocumentGenerator:
    """文档生成器"""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.output_config = self.config.get('output', {})
        self.output_dir = Path(self.output_config.get('directory', './homework_output'))
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def generate(self, assignment: Assignment, content: str,
                       format_type: str = 'docx') -> Path:
        """生成文档"""
        if format_type.lower() == 'docx':
            return await self._generate_docx(assignment, content)
        elif format_type.lower() == 'txt':
            return await self._generate_txt(assignment, content)
        else:
            return await self._generate_docx(assignment, content)

    async def _generate_docx(self, assignment: Assignment, content: str) -> Path:
        """生成Word文档 - 统一英文字体和格式"""
        doc = Document()

        # 设置默认英文字体样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # 标题
        title = doc.add_heading(assignment.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        # 课程信息（英文）
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_run = meta.add_run(
            f"Course: {assignment.course_name}\n"
            f"Due Date: {assignment.due_date.strftime('%Y-%m-%d %H:%M') if assignment.due_date else 'Not Set'}\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        meta_run.font.name = 'Times New Roman'
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        # 分隔线
        doc.add_paragraph()

        # 答案内容 - 统一字体
        content_para = doc.add_paragraph()
        content_run = content_para.add_run(content)
        content_run.font.name = 'Times New Roman'
        content_run.font.size = Pt(12)

        filename = self._generate_filename(assignment, 'docx')
        filepath = self.output_dir / filename

        doc.save(str(filepath))
        return filepath

    async def _generate_txt(self, assignment: Assignment, content: str) -> Path:
        """生成纯文本文件"""
        filename = self._generate_filename(assignment, 'txt')
        filepath = self.output_dir / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"{assignment.title}\n")
            f.write(f"{'=' * 50}\n\n")
            f.write(f"课程：{assignment.course_name}\n")
            f.write(f"截止日期：{assignment.due_date.strftime('%Y-%m-%d %H:%M') if assignment.due_date else '未设置'}\n")
            f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
            f.write(f"{'=' * 50}\n\n")
            f.write(content)

        return filepath

    def _generate_filename(self, assignment: Assignment, extension: str) -> str:
        """生成文件名"""
        template = self.output_config.get('filename_template',
                                         '{course}_{assignment}_{date}')

        course_name = self._sanitize_filename(assignment.course_name)
        assignment_title = self._sanitize_filename(assignment.title)
        date_str = datetime.now().strftime('%Y%m%d_%H%M')

        filename = template.format(
            course=course_name[:20],
            assignment=assignment_title[:30],
            date=date_str
        )

        return f"{filename}.{extension}"

    def _sanitize_filename(self, name: str) -> str:
        """清理文件名中的非法字符"""
        import re
        name = re.sub(r'[<>:"/\\|?*]', '', name)
        name = name.strip()
        return name[:50]

    async def generate_batch(self, assignments: list, contents: list) -> list[Path]:
        """批量生成文档"""
        paths = []
        for assignment, content in zip(assignments, contents):
            path = await self.generate(assignment, content)
            paths.append(path)
        return paths
